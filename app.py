# app.py (Versio 23.0 - Parannettu aggressiivinen TILA C ja dynaaminen puskuri)
import logging
import re
import time
from collections import defaultdict
from io import BytesIO

import docx
import ollama
import streamlit as st
import streamlit.components.v1 as components
from logic import (
    STRATEGIA_SANAKIRJA,
    arvioi_tulokset,
    ehdota_uutta_strategiaa,
    etsi_merkityksen_mukaan,
    etsi_puhtaalla_haulla,
    lataa_resurssit,
    luo_kontekstisidonnainen_avainsana,
    suorita_tarkennushaku,
    tallenna_uusi_strategia,
)
from monitoring import log_performance_stats, setup_performance_logger


def auto_scroll_js():
    """Luo JavaScript-koodin loki-ikkunan automaattiseen vieritykseen."""
    return """
    <script>
        window.setInterval(function() {
            var stLog = parent.document.querySelector('[data-testid="stExpander"] div[data-testid="stVerticalBlock"]');
            if (stLog) {
                stLog.scrollTop = stLog.scrollHeight;
            }
        }, 500);
    </script>
    """


# --- Sivun asetukset ---
st.set_page_config(
    page_title="Raamattu-tutkija v5",
    page_icon="📚",
    layout="wide"
)


# --- Lokituksen Asetukset ---
def setup_logger():
    """Alustaa lokituksen sekä Streamlit-konsoliin että tiedostoon."""
    logger = logging.getLogger()
    if logger.hasHandlers():
        logger.handlers.clear()
    logger.setLevel(logging.INFO)

    timestamp = time.strftime("%Y%m%d-%H%M%S")
    log_filename = f"app_session_{timestamp}.log"
    file_handler = logging.FileHandler(
        log_filename, mode='w', encoding='utf-8'
    )
    file_handler.setFormatter(
        logging.Formatter('%(asctime)s - %(message)s', datefmt="%H:%M:%S")
    )
    logger.addHandler(file_handler)

    return logger


class StreamlitLogHandler(logging.Handler):
    """Ohjaa lokituksen Streamlitin UI-elementtiin."""
    def __init__(self, container):
        super().__init__()
        self.container = container

    def emit(self, record):
        msg = self.format(record)
        if record.levelno == logging.WARNING:
            self.container.warning(msg)
        elif record.levelno >= logging.ERROR:
            self.container.error(msg)
        else:
            self.container.info(msg)


# --- Apufunktiot ---
@st.cache_data(ttl=600)
def hae_asennetut_mallit():
    """Hakee ja palauttaa listan asennetuista Ollama-malleista."""
    try:
        models_data = ollama.list()
        valid_models = [
            model['model'] for model in models_data.get('models', [])
            if 'model' in model
        ]

        if not valid_models:
            st.warning(
                "Yhtään kelvollista Ollama-mallia ei löytynyt. "
                "Tarkista asennuksesi komennolla 'ollama list'."
            )
            return ["Ei löytynyt kelvollisia malleja"]

        return valid_models
    except Exception as e:
        st.error(f"Ollama-yhteys epäonnistui: {e}")
        return ["Yhteyttä ei saatu"]


def lue_syote_data(syote_data):
    """Jäsentää syötetiedon vankasti ja valmistelee sen hakua varten."""
    if not syote_data:
        return None, None, None, None

    if hasattr(syote_data, 'getvalue'):
        sisalto = syote_data.getvalue().decode("utf-8")
    else:
        sisalto = str(syote_data)
    sisalto = sisalto.replace('\r\n', '\n')

    paaotsikko_m = re.search(r"^(.*?)\n", sisalto)
    paaotsikko = paaotsikko_m.group(1).strip() if paaotsikko_m else ""

    hakulauseet = {}
    otsikot = {}

    osiot = re.split(r'\n(?=\d\.\s)', sisalto)

    for osio_teksti in osiot:
        osio_teksti = osio_teksti.strip()
        if not osio_teksti:
            continue
        rivit = osio_teksti.split('\n', 1)
        otsikko = rivit[0].strip()
        kuvaus = rivit[1].strip() if len(rivit) > 1 else ""

        osio_match = re.match(r"^([\d\.]+)", otsikko)
        if osio_match:
            osio_nro = osio_match.group(1).strip('.')
            kuvaus_muokattu = kuvaus.replace('\n', ' ')
            haku = f"{otsikko}: {kuvaus_muokattu}"
            hakulauseet[osio_nro] = haku
            otsikot[osio_nro] = otsikko

    sl_match = re.search(
        r"Sisällysluettelo:(.*?)(?=\n\d\.|\Z)", sisalto, re.DOTALL
    )
    sl_teksti = sl_match.group(1).strip() if sl_match else ""
    return paaotsikko, hakulauseet, otsikot, sl_teksti


def luo_raportti_md(sl, jae_kartta, arvosanat):
    """Luo siistin tekstimuotoisen raportin."""
    md = f"# {sl['otsikko']}\n\n"
    if sl['teksti']:
        md += f"## Sisällysluettelo\n\n{sl['teksti']}\n\n"

    sorted_osiot = sorted(
        jae_kartta.items(),
        key=lambda item: [int(p) for p in item[0].split('.')]
    )
    for osio_nro, data in sorted_osiot:
        arvosana_num = arvosanat.get(osio_nro)
        arvosana_str = (
            f"{arvosana_num:.2f}"
            if isinstance(arvosana_num, (float, int)) else "N/A"
        )
        md += (
            f"## {data['otsikko']} "
            f"(Lopullinen arvosana: {arvosana_str}/10)\n\n"
        )
        if data["jakeet"]:
            for jae in data["jakeet"]:
                md += f"- **{jae['viite']}**: \"{jae['teksti']}\"\n"
        else:
            md += "*Ei jakeita tähän osioon.*\n"
        md += "\n"
    return md


def luo_raportti_doc(sl, jae_kartta, arvosanat):
    """Luo ladattavan Word-dokumentin."""
    doc = docx.Document()
    doc.add_heading(sl['otsikko'], 0)
    if sl['teksti']:
        doc.add_heading("Sisällysluettelo", 1)
        doc.add_paragraph(sl['teksti'])

    sorted_osiot = sorted(
        jae_kartta.items(),
        key=lambda item: [int(p) for p in item[0].split('.')]
    )
    for osio_nro, data in sorted_osiot:
        arvosana_num = arvosanat.get(osio_nro)
        arvosana_str = (
            f"{arvosana_num:.2f}"
            if isinstance(arvosana_num, (float, int)) else "N/A"
        )
        doc.add_heading(
            f"{data['otsikko']} (Lopullinen arvosana: {arvosana_str}/10)", 1
        )
        if data["jakeet"]:
            for jae in data["jakeet"]:
                p = doc.add_paragraph()
                p.add_run(f"{jae['viite']}: ").bold = True
                p.add_run(f"\"{jae['teksti']}\"")
        else:
            doc.add_paragraph("Ei jakeita tähän osioon.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Streamlit-käyttöliittymä ---
st.title("📚 Raamattu-tutkija v5 (Asiantuntija-asetuksin)")

with st.spinner("Valmistellaan hakukonetta..."):
    lataa_resurssit()

if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'tehostesanat' not in st.session_state:
    st.session_state.tehostesanat = {}

col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("Syötä tutkielman aihe ja rakenne")
    uploaded_file = st.file_uploader("1. Lataa .txt-tiedosto", type=['txt'])
    syote_alue = st.text_area(
        "2. Kirjoita tai liitä sisältö tähän", height=350
    )

with col2:
    st.subheader("Asetukset")
    top_k_valinta = st.slider("Jakeita per osio:", 1, 100, 15)
    oppiminen_paalla = st.checkbox(
        "Oppiminen päällä (tallentaa strategiat pysyvästi)",
        value=False,
        help=(
            "Jos tämä on valittuna, sovellus muokkaa `logic.py`-tiedostoa, "
            "kun se löytää parannuksia TILA B:n strategiaparannuksessa."
        )
    )
    with st.expander("Asiantuntija-asetukset"):
        asennetut_mallit = hae_asennetut_mallit()

        default_model = "raamattu-tutkija-model:q4"
        if asennetut_mallit and default_model not in asennetut_mallit:
            default_model = asennetut_mallit[0] if asennetut_mallit else None

        valittu_malli = st.selectbox(
            "Valitse arviointimalli:",
            options=asennetut_mallit,
            index=asennetut_mallit.index(default_model) if default_model and default_model in asennetut_mallit else 0,
            help="Valitse Ollamaan asennettu malli tulosten arviointiin."
        )
        ydinjakeiden_minimi = st.number_input(
            "Ydinjakeiden minimimäärä (TILA A):",
            min_value=2, max_value=10, value=3, step=1,
            help=(
                "Minimimäärä keskiarvoa parempia jakeita, joka vaaditaan "
                "tarkennushaun käynnistämiseksi."
            )
        )
        aggressiivisuus_kerroin = st.slider(
            "Tarkennushaun aggressiivisuus (TILA A):",
            min_value=1, max_value=10, value=3,
            help="Kuinka monta uutta ehdokasta haetaan per heikko jae (esim. 3x)."
        )
        # KORJATTU: Maksimiarvo laskettu 9.50:een
        laatutavoite = st.slider(
            "Laatutavoite (TILA C):",
            min_value=7.0, max_value=9.5, value=8.5, step=0.1,
            key='laatutavoite_slider',
            help=(
                "Jos osion laatu jää tämän alle, käynnistetään iteratiivinen "
                "parannus."
            )
        )
        maksimi_iteraatiot = st.number_input(
            "Maksimi-iteraatiot (TILA C):",
            min_value=1, max_value=10, value=5, step=1,
            help="Kuinka monta kertaa TILA C yrittää saavuttaa laatutavoitteen."
        )
        pakota_tila_c = st.checkbox(
            "Pakota aggressiivinen parannus (TILA C)",
            key='pakota_tila_c_checkbox',
            value=False,
            help=(
                "Jos valittuna, ohjelma ohittaa 'Maksimi-iteraatiot'-rajan ja "
                "jatkaa parantamista, kunnes tavoite saavutetaan tai laatu ei "
                "enää parane. Myös TILA B ohitetaan tarvittaessa."
            )
        )

# --- Syötteen käsittely ja Avainsana-Tehostin ---
koko_syote = ""
if uploaded_file is not None:
    koko_syote += uploaded_file.getvalue().decode("utf-8")
koko_syote += "\n\n" + syote_alue.strip()
koko_syote = koko_syote.strip()

if koko_syote:
    _, hakulauseet_dialog, otsikot_dialog, _ = lue_syote_data(koko_syote)
    if hakulauseet_dialog:
        with st.expander(
            "Avainsana-Tehostin: Vahvista ja lisää osiokohtaiset sanat",
            expanded=True
        ):
            if 'tehostesanat' not in st.session_state:
                st.session_state.tehostesanat = {}

            sorted_osiot_dialog = sorted(
                hakulauseet_dialog.items(),
                key=lambda item: [int(p) for p in item[0].split('.')]
            )

            for osio_nro, haku in sorted_osiot_dialog:
                otsikko = otsikot_dialog.get(osio_nro, "")
                st.markdown(f"**Osio {osio_nro}:** *{otsikko.split(':')[0]}*")

                with st.form(key=f"form_{osio_nro}"):
                    _, tunnistetut_sanat = etsi_merkityksen_mukaan(
                        haku, otsikko, top_k=0
                    )
                    kaikki_vaihtoehdot = sorted(
                        list(tunnistetut_sanat |
                             st.session_state.tehostesanat.get(osio_nro, set()))
                    )
                    valitut = st.multiselect(
                        f"Valitse tehostettavat sanat osiolle {osio_nro}:",
                        options=kaikki_vaihtoehdot,
                        default=st.session_state.tehostesanat.get(
                            osio_nro, list(tunnistetut_sanat)
                        ),
                        key=f"multiselect_{osio_nro}"
                    )
                    st.session_state.tehostesanat[osio_nro] = set(valitut)
                    uusi_sana = st.text_input(
                        f"Lisää uusi sana tai termi osiolle {osio_nro}",
                        key=f"text_input_{osio_nro}"
                    )
                    submitted = st.form_submit_button("Lisää sana")
                    if submitted and uusi_sana:
                        st.session_state.tehostesanat.setdefault(
                            osio_nro, set()
                        ).add(uusi_sana.lower())
                        st.rerun()

            st.info(
                "Voit nyt sulkea tämän laatikon. Asetukset on tallennettu."
            )

suorita_nappi = st.button("Suorita älykäs haku", type="primary")

st.markdown("---")
st.subheader("Prosessi ja tulokset")

if suorita_nappi:
    if not koko_syote:
        st.warning("Syötä aihe ja rakenne.")
    else:
        logger = setup_logger()
        log_container = st.expander("Näytä prosessin loki", expanded=True)
        components.html(auto_scroll_js(), height=0)
        streamlit_handler = StreamlitLogHandler(log_container)
        streamlit_handler.setFormatter(logging.Formatter('%(message)s'))
        logger.addHandler(streamlit_handler)

        perf_writer, perf_file = setup_performance_logger()

        try:
            log_performance_stats(perf_writer, perf_file)
            st.session_state.processing_complete = False
            (
                paaotsikko, hakulauseet,
                otsikot, sl_teksti
            ) = lue_syote_data(koko_syote)
            sl = {"otsikko": paaotsikko, "teksti": sl_teksti}
            jae_kartta = defaultdict(lambda: {"jakeet": [], "otsikko": ""})
            lopulliset_arvosanat = {}

            sorted_hakulauseet = sorted(
                hakulauseet.items(),
                key=lambda item: [int(p) for p in item[0].split('.')]
            )

            ui_laatutavoite = st.session_state.get('laatutavoite_slider', 8.5)
            ui_pakota_tila_c = st.session_state.get('pakota_tila_c_checkbox', False)
            ui_maksimi_iteraatiot = maksimi_iteraatiot

            for i, (osio_nro, haku) in enumerate(sorted_hakulauseet):
                log_container.markdown(
                    f"--- \n ### Käsitellään osiota "
                    f"{i+1}/{len(sorted_hakulauseet)}: "
                    f"{otsikot.get(osio_nro, '')}"
                )

                musta_lista_viitteet = set()
                log_performance_stats(perf_writer, perf_file)

                valitut_tehostesanat_osiolle = (
                    st.session_state.tehostesanat.get(osio_nro, set())
                )
                tulokset, _ = etsi_merkityksen_mukaan(
                    haku,
                    otsikot.get(osio_nro, ''),
                    top_k=top_k_valinta,
                    valitut_tehostesanat=valitut_tehostesanat_osiolle
                )

                musta_lista_viitteet.update(t['viite'] for t in tulokset)
                log_performance_stats(perf_writer, perf_file)

                arvio = arvioi_tulokset(
                    haku, tulokset, malli_nimi=valittu_malli
                )
                log_performance_stats(perf_writer, perf_file)

                if not arvio.get("jae_arviot"):
                    logging.error(
                        "KRIITTINEN: Arviointi epäonnistui. Suoritetaan puhdas haku."
                    )
                    final_tulokset = etsi_puhtaalla_haulla(
                        haku, top_k=top_k_valinta
                    )
                    arvio = arvioi_tulokset(haku, final_tulokset, malli_nimi=valittu_malli)
                    for jae in final_tulokset:
                        vastaava = next((a for a in arvio.get("jae_arviot", []) if a.get('viite') == jae['viite']), None)
                        if vastaava:
                            jae.update(vastaava)
                else:
                    arvioidut_tulokset = []
                    for jae in tulokset:
                        vastaava = next(
                            (a for a in arvio.get("jae_arviot", [])
                             if a.get('viite') == jae['viite']), None
                        )
                        if vastaava:
                            jae.update(vastaava)
                            arvioidut_tulokset.append(jae)
                    final_tulokset = arvioidut_tulokset

                valid_scores = [t.get('arvosana') for t in final_tulokset if t.get('arvosana') is not None]
                alkuperainen_keskiarvo = sum(valid_scores) / len(valid_scores) if valid_scores else 0.0
                lopputulos_keskiarvo = alkuperainen_keskiarvo
                logging.info(f"Alkuperäinen laatuarvio: {alkuperainen_keskiarvo:.2f}/10")

                log_container.markdown("--- \n #### Taso 2: Dynaaminen Parannusalgoritmi")

                dynaaminen_raja_arvo = alkuperainen_keskiarvo
                ydinjakeet = [t for t in final_tulokset if t.get('arvosana', 0) >= dynaaminen_raja_arvo]

                tila_a_ehdot_taynna = len(ydinjakeet) >= ydinjakeiden_minimi

                if tila_a_ehdot_taynna and not ui_pakota_tila_c:
                    logging.info(f"TILA A: Ydinjakeita löytyi {len(ydinjakeet)} kpl. Suoritetaan tarkennushaku.")
                    heikot = sorted([t for t in final_tulokset if t.get('arvosana', 0) < dynaaminen_raja_arvo], key=lambda x: x.get('arvosana', 0))
                    haettava_maara = max(10, min(50, len(heikot) * aggressiivisuus_kerroin))
                    uudet_ehdokkaat = suorita_tarkennushaku(ydinjakeet, musta_lista_viitteet, haettava_maara)
                    if uudet_ehdokkaat:
                        musta_lista_viitteet.update(t['viite'] for t in uudet_ehdokkaat)
                        uudet_arviot = arvioi_tulokset(haku, uudet_ehdokkaat, malli_nimi=valittu_malli).get("jae_arviot", [])
                        for jae in uudet_ehdokkaat:
                            vastaava = next((a for a in uudet_arviot if a.get('viite') == jae['viite']), None)
                            if vastaava:
                                jae.update(vastaava)
                        
                        uudet_parhaat = sorted([j for j in uudet_ehdokkaat if 'arvosana' in j], key=lambda x: x.get('arvosana', 0), reverse=True)
                        korvaus_laskuri = 0
                        for i_korv in range(len(heikot)):
                            if i_korv < len(uudet_parhaat) and uudet_parhaat[i_korv].get('arvosana', 0) > heikot[i_korv].get('arvosana', 0):
                                for idx, item in enumerate(final_tulokset):
                                    if item['viite'] == heikot[i_korv]['viite']:
                                        final_tulokset[idx] = uudet_parhaat[i_korv]
                                        break
                                korvaus_laskuri += 1
                        logging.info(f"Laadunvalvonta valmis. {korvaus_laskuri} jaetta korvattu.")

                elif not tila_a_ehdot_taynna and not ui_pakota_tila_c:
                    logging.warning(f"TILA B: Ydinjakeita ei tarpeeksi ({len(ydinjakeet)}/{ydinjakeiden_minimi}). Siirrytään strategian parannukseen.")
                    arvio_obj = {"kokonaisarvosana": alkuperainen_keskiarvo, "jae_arviot": final_tulokset}
                    ehdotus = ehdota_uutta_strategiaa(haku, arvio_obj)
                    if "virhe" not in ehdotus and ehdotus.get("selite"):
                        avainsanat, selite = ehdotus.get('avainsanat', []), ehdotus.get('selite', '')
                        if oppiminen_paalla:
                            uniikit_sanat = [luo_kontekstisidonnainen_avainsana(s, selite) if s.lower() in STRATEGIA_SANAKIRJA else s for s in avainsanat]
                            tallenna_uusi_strategia(uniikit_sanat, selite)
                        heikot_lkm = len([t for t in final_tulokset if t.get('arvosana', 0) < dynaaminen_raja_arvo])
                        if heikot_lkm > 0:
                            paikkaushaku, _ = etsi_merkityksen_mukaan(haku, otsikot.get(osio_nro, ''), top_k=heikot_lkm, custom_strategiat={s.lower(): selite for s in avainsanat})
                            if paikkaushaku:
                                final_tulokset_hyvat = [t for t in final_tulokset if t.get('arvosana', 0) >= dynaaminen_raja_arvo]
                                final_tulokset = final_tulokset_hyvat + paikkaushaku

                valid_scores_after_ab = [t.get('arvosana') for t in final_tulokset if t.get('arvosana') is not None]
                lopputulos_keskiarvo = sum(valid_scores_after_ab) / len(valid_scores_after_ab) if valid_scores_after_ab else 0.0

                if lopputulos_keskiarvo < ui_laatutavoite or (ui_pakota_tila_c and not tila_a_ehdot_taynna):
                    if ui_pakota_tila_c and not tila_a_ehdot_taynna:
                        logging.warning("TILA C (PAKOTETTU): Ohitetaan TILA B ja siirrytään suoraan iteratiiviseen parannukseen.")
                    
                    log_container.markdown(f"--- \n #### Taso 3: Laadun Tavoittelu (Tavoite: {ui_laatutavoite:.2f})")
                    logging.info(f"LÄHTÖTILANNE TILA C: Keskiarvo {lopputulos_keskiarvo:.2f}/{ui_laatutavoite:.2f}")
                    
                    # KORJATTU: Aggressiivinen tila ohittaa iteraatiorajan
                    iteraatioraja = 100 if ui_pakota_tila_c else ui_maksimi_iteraatiot
                    iteraatio = 0
                    while (lopputulos_keskiarvo < ui_laatutavoite and iteraatio < iteraatioraja):
                        iteraatio += 1
                        logging.info(f"Käynnistetään TILA C -parannusyritys {iteraatio}...")

                        edellinen_keskiarvo = lopputulos_keskiarvo
                        
                        nykyinen_summa = sum(t.get('arvosana', 0) for t in final_tulokset if t.get('arvosana') is not None)
                        tarvittava_summa = ui_laatutavoite * len(final_tulokset)
                        summan_ero = max(0, tarvittava_summa - nykyinen_summa)
                        
                        heikot_c = sorted([t for t in final_tulokset if t.get('arvosana', 0) < lopputulos_keskiarvo], key=lambda x: x.get('arvosana', 0))
                        if not heikot_c:
                            logging.warning("TILA C: Ei enää heikkoja jakeita parannettavaksi. Silmukka päättyy.")
                            break
                        
                        korvattavia_lkm = 0
                        potentiaalinen_nousu = 0
                        for heikko_jae in heikot_c:
                            potentiaalinen_nousu += (10.0 - heikko_jae.get('arvosana', 0))
                            korvattavia_lkm += 1
                            if potentiaalinen_nousu >= summan_ero:
                                break
                        
                        # KORJATTU: Dynaaminen puskuri
                        puskuri = max(3, korvattavia_lkm // 2)
                        haettava_maara_c = korvattavia_lkm + puskuri
                        logging.info(f"Tavoitteeseen vaaditaan {summan_ero:.2f}p. Yritetään korvata {korvattavia_lkm} jaetta hakemalla {haettava_maara_c} uutta ehdokasta.")
                        
                        ydinjakeet_c = [t for t in final_tulokset if t.get('arvosana', 0) >= lopputulos_keskiarvo]
                        if not ydinjakeet_c:
                            ydinjakeet_c = sorted(final_tulokset, key=lambda x: x.get('arvosana', 0), reverse=True)[:5]

                        uudet_ehdokkaat_c = suorita_tarkennushaku(ydinjakeet_c, musta_lista_viitteet, haettava_maara_c)
                        if not uudet_ehdokkaat_c:
                            logging.warning("TILA C: Tarkennushaku ei löytänyt enempää uniikkeja jakeita. Silmukka päättyy.")
                            break
                        musta_lista_viitteet.update(t['viite'] for t in uudet_ehdokkaat_c)

                        logging.info(f"Arvioidaan {len(uudet_ehdokkaat_c)} uutta ehdokasta kerralla...")
                        uudet_arviot_c = arvioi_tulokset(haku, uudet_ehdokkaat_c, malli_nimi=valittu_malli).get("jae_arviot", [])
                        for jae in uudet_ehdokkaat_c:
                            vastaava = next((a for a in uudet_arviot_c if a.get('viite') == jae['viite']), None)
                            if vastaava:
                                jae.update(vastaava)

                        uudet_parhaat_c = sorted([j for j in uudet_ehdokkaat_c if 'arvosana' in j], key=lambda x: x.get('arvosana', 0), reverse=True)
                        
                        korvaus_laskuri_c = 0
                        heikot_c_uudelleen = sorted([t for t in final_tulokset if t.get('arvosana', 0) < edellinen_keskiarvo], key=lambda x: x.get('arvosana', 0))

                        for i_korv in range(len(heikot_c_uudelleen)):
                            if i_korv < len(uudet_parhaat_c):
                                vanha_jae = heikot_c_uudelleen[i_korv]
                                uusi_jae = uudet_parhaat_c[i_korv]
                                if uusi_jae.get('arvosana', 0) > vanha_jae.get('arvosana', 0):
                                    for idx, item in enumerate(final_tulokset):
                                        if item['viite'] == vanha_jae['viite']:
                                            final_tulokset[idx] = uusi_jae
                                            korvaus_laskuri_c += 1
                                            break
                        
                        valid_scores_now = [t.get('arvosana') for t in final_tulokset if t.get('arvosana') is not None]
                        lopputulos_keskiarvo = sum(valid_scores_now) / len(valid_scores_now) if valid_scores_now else 0.0

                        logging.info(f"TILA C: Kierros {iteraatio} valmis. {korvaus_laskuri_c} jaetta korvattu.")
                        if lopputulos_keskiarvo > edellinen_keskiarvo:
                            logging.info(f"TILA C: LAATU PARANI: {edellinen_keskiarvo:.2f} -> {lopputulos_keskiarvo:.2f} ✅")
                        else:
                            logging.warning("TILA C: Laatu ei parantunut tällä kierroksella. Silmukka päättyy.")
                            lopputulos_keskiarvo = edellinen_keskiarvo
                            break
                    
                    if lopputulos_keskiarvo < ui_laatutavoite:
                        logging.warning(f"TILA C: Laatutavoitetta ({ui_laatutavoite:.2f}) ei saavutettu. Lopullinen laatu: {lopputulos_keskiarvo:.2f}/10.")

                jae_kartta[osio_nro]["jakeet"] = sorted(
                    final_tulokset, key=lambda x: x.get('arvosana', 0),
                    reverse=True
                )
                jae_kartta[osio_nro]["otsikko"] = otsikot.get(
                    osio_nro, haku.split(':')[0]
                )
                lopulliset_arvosanat[osio_nro] = lopputulos_keskiarvo

            log_performance_stats(perf_writer, perf_file)

        finally:
            if perf_file:
                perf_file.close()

            logger.removeHandler(streamlit_handler)
            st.session_state.final_report_md = luo_raportti_md(
                sl, jae_kartta, lopulliset_arvosanat
            )
            st.session_state.final_report_doc = luo_raportti_doc(
                sl, jae_kartta, lopulliset_arvosanat
            )
            st.session_state.processing_complete = True
            st.success("Haku suoritettu onnistuneesti!")
            time.sleep(10)
            st.rerun()

if st.session_state.processing_complete:
    st.markdown("---")
    st.subheader("Lopullinen tutkielma")
    st.markdown(st.session_state.final_report_md)
    st.divider()
    st.subheader("Lataa raportti")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Lataa .txt-tiedostona",
            st.session_state.final_report_md,
            "raamattu_tutkielma.txt"
        )
    with col2:
        st.download_button(
            "Lataa .docx-tiedostona",
            st.session_state.final_report_doc,
            "raamattu_tutkielma.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )